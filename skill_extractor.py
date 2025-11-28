# app.py
import streamlit as st
import os
import re
import json
import time
import tempfile
from typing import Tuple, List, Optional

st.markdown("""
<style>

html, body, [data-testid="stAppViewContainer"] {
    background-color: #0e1117 !important;
    color: #ffffff !important;
}

[data-testid="stHeader"] {
    background-color: #0e1117 !important;
}

[data-testid="stSidebar"] {
    background-color: #0e1117 !important;
    color: #ffffff !important;
}

h1, h2, h3, h4, h5, h6, p, span, div, label {
    color: #ffffff !important;
}

/* force dark scrollbars */
::-webkit-scrollbar {
    width: 8px;
}
::-webkit-scrollbar-track {
    background: #1a1e24;
}
::-webkit-scrollbar-thumb {
    background: #3a3f47;
    border-radius: 10px;
}
::-webkit-scrollbar-thumb:hover {
    background: #4d535c;
}

</style>
""", unsafe_allow_html=True)


# ---------------- CSS (chips + uniform styling) ----------------
st.markdown(
    """
    <style>
    .skills-container { margin-bottom: 14px; }
    .skills-row { display:flex; flex-wrap:wrap; gap:8px; margin:8px 0 18px 0; }
    .skill-chip {
        padding:6px 12px;
        border-radius:999px;
        color:#fff;
        font-weight:600;
        font-size:14px;
        font-family: "Inter", "Arial", sans-serif;
        box-shadow: 0 1px 2px rgba(0,0,0,0.15);
    }
    .cat-title { font-weight:700; margin:6px 0 4px 0; color: #e6eef8; }

    /* category colors */
    .lang { background: #1f77b4; }
    .tools { background: #2ca02c; }
    .protocols { background: #ff7f0e; color:#111; }
    .platforms { background: #9467bd; }
    .drivers { background: #d62728; }
    .other { background: #7f8c8d; }
    /* Highlight header row */
    .table-header {
        font-size: 20px !important;
        font-weight: 800 !important;
        color: #ffffff !important;
        padding: 8px 0px;
    }

    /* Highlight resume name */
    .resume-name {
        font-size: 17px !important;
        font-weight: 700 !important;
        color: #00b4ff !important;     /* light-blue highlight */
        padding-top: 6px;
    }

    /* Optional: stronger highlight */
    .resume-name:hover {
        color: #33c9ff !important;
        transition: 0.2s;
    }

    .section-title { font-weight:800; margin-top:10px; margin-bottom:6px; font-size:18px; }
    /* style Start button */
    div.stButton > button {
        background-color:#2E8B57;
        color:white;
        font-size:15px;
        padding:10px 18px;
        border-radius:8px;
    }
    div.stButton > button:hover {
        background-color:#3CB371;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ---------- Optional dependencies used by original script ----------
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx
except Exception:
    docx = None

# ---------- Optional Gemini client wrapper ----------
GEMINI_CLIENT = None
try:
    from google import genai as _genai
    GEMINI_CLIENT = _genai
except Exception:
    GEMINI_CLIENT = None

# ---------- KEY RETRIEVAL ----------
def get_gemini_key() -> str:
    try:
        if "GEMINI_API_KEY" in st.secrets:
            return st.secrets["GEMINI_API_KEY"]
    except Exception:
        pass
    if os.getenv("GEMINI_API_KEY"):
        return os.getenv("GEMINI_API_KEY")
    return ""

# ---------- Text extraction ----------
def extract_text_from_pdf(path: str) -> str:
    if pdfplumber is None:
        raise RuntimeError("Install pdfplumber: pip install pdfplumber")
    parts = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            parts.append(p.extract_text() or "")
    return re.sub(r'\s+', ' ', " ".join(parts)).strip()

def extract_text_from_docx(path: str) -> str:
    if docx is None:
        raise RuntimeError("Install python-docx: pip install python-docx")
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    return re.sub(r'\s+', ' ', " ".join(parts)).strip()

def extract_text_from_txt(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    return re.sub(r'\s+', ' ', raw.decode('utf-8', errors='ignore')).strip()

def extract_text(path: str) -> str:
    p = path.lower()
    if p.endswith(".pdf"):
        return extract_text_from_pdf(path)
    if p.endswith(".docx"):
        return extract_text_from_docx(path)
    if p.endswith(".txt"):
        return extract_text_from_txt(path)
    raise ValueError("Unsupported file type. Supported: .pdf, .docx, .txt")

# ---------- Keyword list + normalization + categorization ----------
BASE_KEYWORDS = [
    "c", "c++", "c#", "python", "java", "javascript", "typescript", "go", "rust", "ruby", "php", "scala", "kotlin", "swift", "r",
    "react", "angular", "vue", "next.js", "svelte", "html", "css", "sass", "tailwind",
    "node.js", "express", "django", "flask", "spring boot", "spring", "laravel", "asp.net",
    "sql", "postgresql", "mysql", "mongodb", "redis", "oracle", "mssql", "cassandra",
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible", "helm",
    "jenkins", "github actions", "gitlab-ci", "circleci",
    "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "xgboost", "lightgbm", "nlp", "opencv", "spacy",
    "spark", "hadoop", "etl", "airflow",
    "embedded linux", "yocto", "petalinux", "u-boot", "device tree", "kernel", "linux kernel", "bsp",
    "arm", "raspberry pi", "stm32", "nxp", "imx", "qualcomm",
    "i2c", "spi", "uart", "gpio", "pcie", "usb", "ethernet", "can", "i2s",
    "board bring-up", "firmware", "bootloader", "driver development", "kernel drivers", "device drivers",
    "git", "gdb", "cmake", "make", "gcc", "clang", "vivado", "quartus", "jtag",
    "linux", "bash", "shell", "systemd", "sysvinit", "excel", "tableau", "power bi", "docker-compose"
]

NORMALIZE_MAP = {
    "react.js": "react",
    "reactjs": "react",
    "nodejs": "node.js",
    "node js": "node.js",
    "powerbi": "power bi",
    "u boot": "u-boot",
    "u_boot": "u-boot",
    "device-tree": "device tree",
    "embedded c": "c",
    "c plus plus": "c++",
    "cplusplus": "c++",
    "usb 3 0": "usb 3.0",
    "usb3.0": "usb 3.0",
    "wi fi": "wi-fi",
    "wi-fi": "wi-fi",
    "i 2 c": "i2c",
    "i 2 s": "i2s",
    "yocto project": "yocto",
    "petalinux sdk": "petalinux",
    "system verilog": "systemverilog",
    "devops": "devops",
    "microcontrollers": "microcontroller"
}

def normalize_token(tok: str) -> str:
    if not tok:
        return tok
    s = tok.strip().lower()
    s = re.sub(r'[\_\t]+', ' ', s)
    s = re.sub(r'\s*[\/\\]+\s*', ' / ', s)
    s = s.replace(',', ' ')
    s = re.sub(r'\s+', ' ', s)
    for k, v in NORMALIZE_MAP.items():
        if s == k or k in s:
            s = s.replace(k, v)
    s = re.sub(r'(?<!\d)\.(?!\d)', ' ', s)
    s = s.strip()
    s = re.sub(r'\busb\s+3\s+0\b', 'usb 3.0', s)
    s = re.sub(r'\s+', ' ', s)
    if s == 'node':
        s = 'node.js'
    if s == 'reactjs':
        s = 'react'
    return s

def dedupe_preserve_order(items: List[str]) -> List[str]:
    seen = set()
    out = []
    for x in items:
        if x not in seen:
            seen.add(x)
            out.append(x)
    return out

LANGUAGES = {"c", "c++", "c#", "python", "java", "javascript", "typescript", "go", "rust", "ruby", "php", "scala", "kotlin", "swift", "r"}
TOOLS = {"git", "gdb", "cmake", "make", "gcc", "clang", "vivado", "quartus", "jtag", "docker", "helm", "ansible"}
PROTOCOLS = {"i2c", "spi", "uart", "gpio", "pcie", "usb", "ethernet", "can", "i2s", "wi-fi", "wifi", "lte", "bluetooth"}
PLATFORMS = {"embedded linux", "yocto", "petalinux", "u-boot", "raspberry pi", "stm32", "arm", "nxp", "imx", "xilinx zynq", "xilinx rfsoc", "xilinx mpsoc"}
DRIVERS = {"kernel drivers", "device drivers", "driver development", "bootloader", "board bring-up", "bsp", "firmware", "kernel", "linux kernel"}
OTHER_HINTS = {"linux", "bash", "shell", "systemd", "sysvinit", "excel", "tableau", "power bi", "etl", "spark", "hadoop"}

def categorize_skill(skill: str) -> str:
    low = skill.lower()
    if low in LANGUAGES:
        return "languages"
    if low in TOOLS:
        return "tools"
    if low in PROTOCOLS:
        return "protocols"
    if low in PLATFORMS:
        return "platforms"
    if low in DRIVERS:
        return "drivers"
    if low in OTHER_HINTS:
        return "other"
    if any(k in low for k in ("driver", "kernel", "bootloader", "bsp", "board bring-up", "firmware")):
        return "drivers"
    if any(k in low for k in ("linux", "embedded", "yocto", "petalinux", "u-boot")):
        return "platforms"
    if any(k in low for k in ("aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible")):
        return "tools"
    if any(k in low for k in ("i2c", "spi", "uart", "gpio", "usb", "ethernet", "can", "i2s", "bluetooth", "wi-fi")):
        return "protocols"
    if any(k in low for k in ("python", "java", "c++", "c#", "javascript", "typescript", "go", "rust")):
        return "languages"
    return "other"

# ---------- Gemini prompt builder & caller ----------
def build_skills_prompt(resume_text: str, max_chars=15000) -> str:
    if len(resume_text) > max_chars:
        resume_text = resume_text[:max_chars]
    prompt = f'''
You are an extractor. Given the resume text below, return ONLY a single JSON object:

{{"skills": [<list of canonical short skill strings>] }}

Rules:
- Return skill tokens like "python", "c++", "embedded linux", "device tree", "u-boot", "yocto", "i2c", "spi", "git".
- Normalize common variants (react.js -> react, node js -> node.js, powerbi -> power bi).
- Deduplicate and return only skills actually mentioned in the resume.
- Do NOT include company names, addresses, or long descriptive sentences.
- Output EXACTLY one JSON object and nothing else.

Resume:
\"\"\"{resume_text}\"\"\"
'''.strip()
    return prompt

def call_gemini_for_skills(resume_text: str, api_key: str, max_retries=2) -> Tuple[Optional[List[str]], str]:
    if not api_key or GEMINI_CLIENT is None:
        return None, ""
    prompt = build_skills_prompt(resume_text)
    try:
        client = GEMINI_CLIENT.Client(api_key=api_key)
    except Exception:
        try:
            GEMINI_CLIENT.configure(api_key=api_key)
            client = GEMINI_CLIENT
        except Exception:
            return None, ""
    for attempt in range(max_retries + 1):
        try:
            resp = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
            raw = resp.text if hasattr(resp, "text") else str(resp)
            m = re.search(r'\{.*\}', raw, flags=re.DOTALL)
            if m:
                try:
                    obj = json.loads(m.group(0))
                    skills = obj.get("skills", [])
                    if isinstance(skills, list):
                        return skills, raw
                except Exception:
                    pass
            arr_match = re.search(r'\[\s*([^\]]+?)\s*\]', raw, flags=re.DOTALL)
            if arr_match:
                items = re.findall(r'["\']([^"\']+)["\']', arr_match.group(0))
                if items:
                    return items, raw
            return None, raw
        except Exception as e:
            if attempt < max_retries:
                time.sleep(0.6*(attempt+1))
                continue
            return None, "<error: {}>".format(e)

# ---------- Processing a single file (resume or JD) ----------
def process_resume_file(path: str, api_key: str):
    text = extract_text(path)
    skills_raw, _raw_model = call_gemini_for_skills(text, api_key)
    if skills_raw is None:
        found = []
        text_low = text.lower()
        for kw in BASE_KEYWORDS:
            if re.search(r'\b' + re.escape(kw) + r'\b', text_low):
                found.append(kw)
        skills_raw = found
    normalized = []
    for s in skills_raw:
        if not isinstance(s, str):
            continue
        tok = normalize_token(s)
        if not tok:
            continue
        normalized.append(tok)
    normalized = dedupe_preserve_order(normalized)
    final = []
    for s in normalized:
        s2 = s.strip()
        for k, v in NORMALIZE_MAP.items():
            if s2 == k or k in s2:
                s2 = s2.replace(k, v)
        s2 = re.sub(r'\s+', ' ', s2).strip()
        final.append(s2)
    final = dedupe_preserve_order(final)
    categories = {"languages": [], "tools": [], "protocols": [], "platforms": [], "drivers": [], "other": []}
    for s in final:
        cat = categorize_skill(s)
        categories[cat].append(s)
    return {
        "all_skills": final,
        "categories": categories,
    }

# ---------- Streamlit UI ----------
st.set_page_config(page_title="JD vs Resumes — Skill Matcher", layout="wide")
st.title("JD vs Resumes — Skill Matcher")

st.markdown(
    """
Upload a **single Job Description (JD)** file (pdf, docx, txt) and one or more **resumes** (pdf, docx, txt).
Click **Start Searching** to extract skills from the JD and each resume, then show Matched and Missing skills per resume.
"""
)

api_key = get_gemini_key()
if api_key:
    st.success("Gemini API key found (will attempt model extraction).")
else:
    st.info("No Gemini API key found — falling back to local keyword scanner.")

# --- Two-column layout for JD + Resume Uploads ---
col1, col2 = st.columns(2)

with col1:
    st.subheader("Job Description (JD)")
    jd_file = st.file_uploader("Upload JD", type=["pdf", "docx", "txt"], accept_multiple_files=False)

with col2:
    st.subheader("Resumes")
    resume_files = st.file_uploader("Upload Resumes", type=["pdf", "docx", "txt"], accept_multiple_files=True)

# Start button
start = st.button("🚀 Start Searching", use_container_width=True)

# When the button is pressed we run processing
if start:
    if jd_file is None:
        st.warning("Please upload a Job Description (JD) file.")
    elif not resume_files:
        st.warning("Please upload at least one Resume.")
    else:
        # ---------- JD processing ----------
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(jd_file.name)[1]) as tmp:
            tmp.write(jd_file.read())
            jd_tmp = tmp.name

        try:
            jd_out = process_resume_file(jd_tmp, api_key)
        except Exception as e:
            st.error(f"Failed to process JD: {e}")
            jd_out = None
        finally:
            try:
                os.remove(jd_tmp)
            except:
                pass

        jd_skills = set(jd_out["all_skills"]) if jd_out else set()
        st.markdown(f"### JD Skills ({len(jd_skills)})")
        if jd_skills:
            # show JD skills as chips too
            st.markdown('<div class="skills-container">', unsafe_allow_html=True)
            display_order = [
                ("languages", "Programming / Languages", "lang"),
                ("tools", "Tools & Devops", "tools"),
                ("protocols", "Protocols & Interfaces", "protocols"),
                ("platforms", "Platforms & Embedded", "platforms"),
                ("drivers", "Drivers / Firmware", "drivers"),
                ("other", "Other", "other"),
            ]
            for key, pretty, css_class in display_order:
                items = jd_out["categories"].get(key, [])
                if not items:
                    continue
                st.markdown(f'<div class="cat-title">{pretty} — {len(items)}</div>', unsafe_allow_html=True)
                chips_html = '<div class="skills-row">'
                for s in items:
                    safe_s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    chips_html += f'<div class="skill-chip {css_class}">{safe_s}</div>'
                chips_html += '</div>'
                st.markdown(chips_html, unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.write("_No skills detected in JD._")

        # ---------- Resume processing ----------
        results = {}
        with st.spinner("Processing resumes..."):
            for f in resume_files:
                suffix = os.path.splitext(f.name)[1].lower()
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(f.read())
                    tmp_path = tmp.name
                try:
                    out = process_resume_file(tmp_path, api_key)
                    results[f.name] = out
                except Exception as e:
                    results[f.name] = {"error": str(e)}
                finally:
                    try:
                        os.remove(tmp_path)
                    except:
                        pass

        # ---------- Display results in table-like rows ----------
# header row
hcol1, hcol2, hcol3, hcol4 = st.columns([1.2, 2, 4, 4])
hcol1.markdown("**Resume**")
hcol2.markdown("**Sections**")
hcol3.markdown("**Matching Skills**")
hcol4.markdown("**Missing Skills**")

# ordering and display metadata
display_order = [
    ("languages", "Programming / Languages", "lang"),
    ("tools", "Tools & Devops", "tools"),
    ("protocols", "Protocols & Interfaces", "protocols"),
    ("platforms", "Platforms & Embedded", "platforms"),
    ("drivers", "Drivers / Firmware", "drivers"),
    ("other", "Other", "other"),
]

def render_sections_column(col, categories_dict):
    """
    Renders the category titles horizontally as small chips in a column `col`.
    """
    # build inline HTML for section chips
    chips = '<div style="display:flex; gap:8px; flex-wrap:wrap; align-items:center;">'
    for key, pretty, css_class in display_order:
        items = categories_dict.get(key, [])
        if not items:
            continue
        # small section chip (slightly smaller than skill chips)
        chips += f'<div class="skill-chip {css_class}" style="padding:4px 8px; font-size:12px; font-weight:700;">{pretty} — {len(items)}</div>'
    chips += '</div>'
    col.markdown(chips, unsafe_allow_html=True)

def render_skills_grouped(col, cat_dict):
    """
    Renders grouped skills (dict keyed by category) into the provided Streamlit column.
    Groups are shown in the same order as display_order.
    """
    html = '<div class="skills-container">'
    for key, pretty, css_class in display_order:
        items = cat_dict.get(key, [])
        if not items:
            continue
        html += f'<div class="cat-title">{pretty} — {len(items)}</div>'
        html += '<div class="skills-row">'
        for s in items:
            safe_s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html += f'<div class="skill-chip {css_class}">{safe_s}</div>'
        html += '</div>'
    html += '</div>'
    col.markdown(html, unsafe_allow_html=True)

# loop over resumes and show one row per resume
for fname, out in results.items():
    # set up four columns per resume row
    col_resume, col_sections, col_matched, col_missing = st.columns([1.2, 2, 4, 4])

    # column 1: resume name
    col_resume.markdown(f"**{fname}**")

    if "error" in out:
        col_matched.error(out["error"])
        # show empty sections / missing
        col_sections.write("")
        col_missing.write("")
        continue

    resume_skills = set(out.get("all_skills", []))
    matched = sorted(list(resume_skills & jd_skills)) if jd_skills else sorted(list(resume_skills))
    missing = sorted(list(jd_skills - resume_skills)) if jd_skills else []

    # categorize lists
    def categorize_list(arr):
        ret = {"languages": [], "tools": [], "protocols": [], "platforms": [], "drivers": [], "other": []}
        for s in arr:
            cat = categorize_skill(s)
            ret[cat].append(s)
        return ret

    matched_cat = categorize_list(matched)
    missing_cat = categorize_list(missing)

    # column 2: sections (show which sections appear in resume or JD)
    # We'll build a combined view: sections present in either JD or resume (for helpful context)
    combined_sections = {"languages": [], "tools": [], "protocols": [], "platforms": [], "drivers": [], "other": []}
    # prefer JD categories if JD provided, else resume categories (so numbers make sense)
    if jd_skills:
        # use jd_out categories presence
        for k in combined_sections.keys():
            combined_sections[k] = jd_out["categories"].get(k, [])
    else:
        # fallback to resume categories
        combined_sections = categorize_list(list(resume_skills))

    # column 3: matched skills grouped by category (chips)
    if matched:
        render_skills_grouped(col_matched, matched_cat)
    else:
        col_matched.write("_No matched skills._")

    # column 4: missing skills grouped by category (chips)
    if missing:
        render_skills_grouped(col_missing, missing_cat)
    else:
        col_missing.write("_No missing JD skills._")

